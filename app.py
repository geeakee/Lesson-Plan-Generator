import streamlit as st
import google.generativeai as genai
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
import json
import io

# --- 1. CONFIGURATION ---
st.set_page_config(page_title="DepEd AI Lesson Planner", layout="wide")

st.title("📚 AI Weekly Lesson Plan Generator (DLL)")
st.markdown("Generates a formatted Word file based on your inputs and module.")

# --- 2. SIDEBAR INPUTS ---
with st.sidebar:
    st.header("1. API Setup")
    api_key = st.text_input("Enter Google Gemini API Key", type="password", help="Get this from aistudio.google.com")
    
    st.header("2. Class Details")
    grade_level = st.text_input("Grade Level", "Grade 5")
    subject = st.text_input("Subject", "Science")
    quarter = st.text_input("Quarter/Week", "Quarter 1 - Week 1")
    
    st.header("3. Standards")
    content_std = st.text_area("Content Standard", height=100)
    perf_std = st.text_area("Performance Standard", height=100)
    competency = st.text_area("Learning Competency", height=100)
    
    st.header("4. Upload Module")
    uploaded_file = st.file_uploader("Upload Module (Text/PDF)", type=['txt', 'pdf'])

# --- 3. DAILY OBJECTIVES INPUT ---
st.subheader("Daily Objectives")
cols = st.columns(5)
days = ["Monday", "Tuesday", "Wednesday", "Thursday", "Friday"]
objectives = {}

for i, day in enumerate(days):
    with cols[i]:
        objectives[day] = st.text_area(f"{day} Objective", height=100)

# --- 4. GENERATION LOGIC ---
if st.button("Generate Lesson Plan", type="primary"):
    if not api_key:
        st.error("Please enter your API Key in the sidebar.")
    elif not uploaded_file:
        st.error("Please upload a module/material.")
    else:
        with st.spinner("AI is reading your module and creating the lesson plan... (This takes about 30 seconds)"):
            try:
                # Configure AI
                genai.configure(api_key=api_key)
                
                # Process File
                # Note: For simplicity in this free script, we treat PDF as text if possible, 
                # or you can copy-paste text. Here we assume text for maximum compatibility 
                # or use Gemini's internal file handling if setup. 
                # To keep it 100% error-free for non-devs, we will extract text if it's simple,
                # otherwise we pass the file to Gemini 1.5 Flash (which can read PDFs).
                
                # Reading the file bytes
                file_bytes = uploaded_file.getvalue()
                
                # Setup Model
                model = genai.GenerativeModel('gemini-1.5-flash')
                
                # Create the Prompt
                prompt_text = f"""
                You are a DepEd Curriculum Expert. Create a Weekly Lesson Plan (DLL) for {grade_level} {subject}.
                
                Context:
                - Content Standard: {content_std}
                - Performance Standard: {perf_std}
                - Competency: {competency}
                - Daily Objectives: {json.dumps(objectives)}
                
                Task:
                Using the attached file as the learning resource, fill out the specific DLL matrix.
                
                Output Format:
                Return ONLY valid JSON. The JSON should be a list of objects where keys are the lesson parts and values are objects containing days.
                
                Structure:
                {{
                    "review": {{"Monday": "...", "Tuesday": "...", ...}},
                    "purpose": {{"Monday": "...", "Tuesday": "...", ...}},
                    "examples": {{"Monday": "...", "Tuesday": "...", ...}},
                    "discuss_1": {{"Monday": "...", "Tuesday": "...", ...}},
                    "discuss_2": {{"Monday": "...", "Tuesday": "...", ...}},
                    "mastery": {{"Monday": "...", "Tuesday": "...", ...}},
                    "application": {{"Monday": "...", "Tuesday": "...", ...}},
                    "generalization": {{"Monday": "...", "Tuesday": "...", ...}},
                    "evaluation": {{"Monday": "...", "Tuesday": "...", ...}},
                    "remediation": {{"Monday": "...", "Tuesday": "...", ...}}
                }}
                """
                
                # Generate content
                # We pass the prompt + the file data (Gemini handles the PDF parsing natively)
                response = model.generate_content([
                    {'mime_type': uploaded_file.type, 'data': file_bytes},
                    prompt_text
                ])
                
                # Parse JSON
                # Clean up json string if AI adds backticks
                json_str = response.text.replace("```json", "").replace("```", "")
                data = json.loads(json_str)
                
                # --- 5. CREATE WORD DOC ---
                doc = Document()
                
                # Title
                style = doc.styles['Normal']
                font = style.font
                font.name = 'Arial'
                font.size = Pt(11)
                
                doc.add_heading(f'Daily Lesson Log - {subject} {grade_level}', 0)
                doc.add_paragraph(f'Week: {quarter}')
                doc.add_paragraph(f'Content Standard: {content_std}')
                doc.add_paragraph(f'Performance Standard: {perf_std}')
                
                # Create Table
                table = doc.add_table(rows=1, cols=6)
                table.style = 'Table Grid'
                
                # Header Row
                hdr_cells = table.rows[0].cells
                hdr_cells[0].text = "Parts of the Lesson"
                for i, day in enumerate(days):
                    hdr_cells[i+1].text = day
                
                # Row Mapping
                row_labels = [
                    ("review", "Reviewing previous lesson"),
                    ("purpose", "Establishing purpose"),
                    ("examples", "Presenting examples"),
                    ("discuss_1", "Discussing new concepts #1"),
                    ("discuss_2", "Discussing new concepts #2"),
                    ("mastery", "Developing Mastery"),
                    ("application", "Practical application"),
                    ("generalization", "Making generalizations"),
                    ("evaluation", "Evaluating learning"),
                    ("remediation", "Additional activities")
                ]
                
                for key, label in row_labels:
                    row_cells = table.add_row().cells
                    row_cells[0].text = label
                    if key in data:
                        for i, day in enumerate(days):
                            # AI might miss a day key, so use .get
                            row_cells[i+1].text = data[key].get(day, "")
                
                # Save to memory buffer
                doc_io = io.BytesIO()
                doc.save(doc_io)
                doc_io.seek(0)
                
                st.success("Lesson Plan Generated!")
                
                # Download Button
                st.download_button(
                    label="Download Word File (.docx)",
                    data=doc_io,
                    file_name="Lesson_Plan.docx",
                    mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
                )

            except Exception as e:
                st.error(f"An error occurred: {e}")
